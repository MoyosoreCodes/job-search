import re
import os
import time
from docx import Document
import openai

class CoverLetterGenerator:
    def __init__(self, api_key):
        self.client = openai.OpenAI(api_key=api_key)

    def generate_multiple(self, jobs_data, resume_data, results_folder):
        cover_letters_folder = os.path.join(results_folder, 'cover_letters')
        os.makedirs(cover_letters_folder, exist_ok=True)
        count = 0
        for job in jobs_data:
            print(f"Generating letter for {job.get('Company')}...")
            cover_letter, error = self._generate_letter(resume_data, job)
            if cover_letter:
                filepath = self._save_letter(cover_letter, job, cover_letters_folder)
                if filepath:
                    job['Cover Letter Generated'] = "Yes"
                    count += 1
                    print(f"Saved: {os.path.basename(filepath)}")
            else:
                print(f"Failed: {error}")
            time.sleep(1)
        return count

    def _generate_letter(self, resume_data, job_data):
        try:
            prompt = (
                "Write a natural, human cover letter. Rules:\n"
                "- Sound like a real person, not AI\n"
                "- No em dashes at all, use regular dashes or commas\n"
                "- Use simple, everyday language\n"
                "- Make minor grammar imperfections but stay professional\n"
                "- No AI phrases like \"I am excited\", \"I would love to\"\n"
                "- Use contractions\n"
                "- Vary sentence length\n"
                "- Under 250 words\n\n"
                f"RESUME INFO:\nSkills: {', '.join(resume_data.get('skills', [])[:8]) or 'Not specified'}\n"
                f"Experience: {' | '.join(resume_data.get('experience', [])[:2]) or 'Not specified'}\n\n"
                f"JOB:\nCompany: {job_data.get('Company', 'Not specified')}\nPosition: {job_data.get('Job Title', 'Not specified')}\n\n"
                'Format: Start with "Dear Hiring Manager," end with "Best regards,"'
            )
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Write natural, human-like cover letters. Avoid AI patterns and corporate speak."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=400,
                temperature=0.7
            )
            letter = response.choices[0].message.content.strip()
            if "—" in letter:
                return None, "Contains em dashes"
            if len([w for w in letter.split() if len(w) > 12]) > 3:
                return None, "Too many complex words"
            return letter, None
        except Exception as e:
            return None, f"Generation error: {str(e)}"

    def _save_letter(self, letter, job_data, folder):
        try:
            doc = Document()
            doc.add_heading(f"Cover Letter - {job_data.get('Company', 'Unknown')}", level=1)
            doc.add_paragraph(f"Position: {job_data.get('Job Title', 'Not specified')}")
            doc.add_paragraph("")
            doc.add_paragraph(letter)
            safe_company = re.sub(r'[^\w\s-]', '', job_data.get('Company', 'Unknown')).strip()
            filename = f"Cover_Letter_{safe_company}.docx"
            filepath = os.path.join(folder, filename)
            doc.save(filepath)
            return filepath
        except Exception as e:
            print(f"Save error: {e}")
            return None
